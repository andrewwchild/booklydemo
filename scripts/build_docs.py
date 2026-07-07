#!/usr/bin/env python3
"""Generate Word doc (demo script + architecture defense) and PDF one-pager."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parent.parent
WORD_OUTPUT = ROOT / "Bookly_CS_Agent_Demo_Script.docx"
PDF_OUTPUT = ROOT / "Bookly_CS_Agent_One_Pager.pdf"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
    doc.add_paragraph()


def build_word_doc() -> None:
    doc = Document()

    # Title
    title = doc.add_heading("Bookly Customer Support Agent", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Live Demo Script & Architecture Defense Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x4F, 0x8C, 0xFF)
    doc.add_paragraph("Andrew Child  ·  github.com/andrewwchild/booklydemo  ·  booklydemo.streamlit.app")
    doc.add_paragraph()

    _add_heading(doc, "Thesis", 1)
    doc.add_paragraph(
        "The best CX agents don't guess — they clarify, then act on real data. "
        "Wrong answers destroy trust faster than one extra question. "
        "The LLM handles natural language; every transactional fact comes from a tool."
    )

    _add_heading(doc, "Pre-Demo Checklist", 1)
    _add_bullets(doc, [
        "Streamlit app open with a fresh conversation (streamlit run app/streamlit_app.py)",
        "Optional backup: terminal ready with python evals/run_evals.py",
        "Deployed URL: booklydemo.streamlit.app",
    ])

    _add_heading(doc, "Demo Reference Data", 1)
    _add_table(doc,
        ["Order", "Email", "Status", "Use For"],
        [
            ["ORD-1001", "alice@example.com", "Shipped (UPS)", "Happy-path order lookup"],
            ["ORD-1006", "frank@example.com", "Delayed", "Empathy + grounded delay reason"],
            ["ORD-1003", "carol@example.com", "Delivered", "Successful refund flow"],
            ["ORD-1004", "david@example.com", "Window expired", "Policy enforcement edge case"],
            ["ORD-1008", "—", "Cancelled", "Edge-case order status"],
        ],
    )
    _add_table(doc,
        ["Book", "Stock"],
        [
            ["Fourth Wing", "Out of stock"],
            ["Atomic Habits", "In stock"],
        ],
    )

    _add_heading(doc, "Opening (30 seconds)", 1)
    doc.add_paragraph(
        '"Bookly is a fictional online bookstore. Support gets the same five questions over and over: '
        "where's my order, can I return this, is this book in stock, what's your policy, and I can't log in."
    )
    doc.add_paragraph(
        "My thesis: the best support agents don't guess — they clarify, then act on real data. "
        "I built a grounded agent with seven tools, identity verification before refunds, "
        "human handoff with conversation summaries, and a 30-case eval harness. Let me show you."
    )

    _add_heading(doc, "Live Demo Flows", 1)

    flows = [
        ("Flow 1: Clarify-First Order Lookup (~90 sec)", [
            'Type: "Where is my order?" → Agent asks for order ID (no hallucination)',
            "Type: ORD-1001 → lookup_order → UPS tracking + delivery estimate",
            'Say: "The LLM handles language. Facts come from tools."',
        ]),
        ("Flow 2: Delayed Shipment (~60 sec)", [
            "New conversation. Type: ORD-1006",
            "Shows delayed status with carrier reason (Weather delay in Memphis hub)",
            'Say: "Grounded data, not generated fluff. Same tool hits OMS in production."',
        ]),
        ("Flow 3: Inventory Check (~60 sec)", [
            'Type: "Is Fourth Wing in stock?" → Out of stock + restock context',
            'Type: "Do you have Atomic Habits?" → In stock with price',
            'Say: "Stock levels need a live API call — you can\'t RAG your way to inventory."',
        ]),
        ("Flow 4: Policy FAQ (~45 sec)", [
            'Type: "What\'s your return policy?" → get_policy → 30-day window',
            'Say: "Policies are RAG-friendly in production; here they\'re canonical structured data."',
        ]),
        ("Flow 5: Multi-Turn Refund + Identity Verification (~2 min) ⭐", [
            'Turn 1: "I want a refund" → asks order ID',
            "Turn 2: ORD-1003 → asks reason (does NOT jump to order lookup)",
            'Turn 3: "The book arrived damaged" → asks email',
            "Turn 4: carol@example.com → verify_customer_identity → initiate_refund → RFD issued",
            'Say: "Verify before act. Email must match. Verification ≠ approval."',
        ]),
        ("Flow 6: Edge Case — Return Window Expired (~60 sec)", [
            "Refund flow for ORD-1004, david@example.com, reason: wrong item",
            "Identity passes; refund rejected — Return window expired",
            'Say: "Agent doesn\'t override business rules."',
        ]),
        ("Flow 7: Human Handoff (~60 sec)", [
            '"Where is my order?" → ORD-1001 → lookup succeeds',
            '"I need to speak to a real person" → escalate_to_human → ticket + summary',
            'Say: "Human agent gets full context — customer never repeats themselves."',
        ]),
        ("Flow 8: Password Reset (optional, 30 sec)", [
            '"I can\'t log in" → asks email → alice@example.com → send_password_reset',
        ]),
    ]
    for title_text, steps in flows:
        _add_heading(doc, title_text, 2)
        _add_bullets(doc, steps)

    _add_heading(doc, "Eval Harness (90 seconds)", 1)
    doc.add_paragraph("Run in terminal: python evals/run_evals.py")
    _add_bullets(doc, [
        "30 golden cases — clarify-first, tool selection, multi-turn refunds, identity, escalation",
        "100% pass rate against rules engine, runs in under 1 second",
        "CI-ready regression safety when prompts or tools change",
    ])

    _add_heading(doc, "Architecture Walkthrough (3 minutes)", 1)
    doc.add_paragraph(
        "Customer → Streamlit UI → BooklyAgent (orchestrator)\n"
        "                              ├── ConversationMemory (history + slots)\n"
        "                              ├── System prompt (clarify-first rules)\n"
        "                              └── OpenAI function calling\n"
        "                                    └── Bookly tools (orders, catalog, policies)"
    )
    _add_bullets(doc, [
        "UI — Streamlit for fast demo; production = chat widget + auth",
        "Orchestrator — Thin LLM ↔ tool loop. No LangChain.",
        "Memory — Per-session slots so multi-turn flows don't re-ask",
        "Tools — 7 functions; swap JSON for real APIs without touching agent",
        "Dual engine — LLM with API key; rules fallback for offline demo + evals",
    ])

    _add_heading(doc, "Architecture Decisions — How to Defend Them", 1)

    decisions = [
        ("1. Tools over RAG for transactions",
         "Order status, refunds, and inventory use function calling — not vector search.",
         "RAG can't answer 'status of ORD-1006 right now.' Refunds are actions, not text.",
         "I'd still use RAG for long-tail FAQ. Core intents are API-shaped."),
        ("2. Clarify-first, not answer-first",
         "'Where's my order?' asks for order ID before any tool call.",
         "Prevents confident wrong answers. Matches how good human agents work.",
         "Production pre-fills from logged-in session. Clarify-first is the safe default for anonymous chat."),
        ("3. Slot memory + regex extraction",
         "ConversationMemory tracks order_id, reason, email, book_title across turns.",
         "Refund needs 3 fields in order. Regex catches ORD-XXXX and emails reliably.",
         "Production uses structured outputs. The state-machine pattern for high-risk flows stays."),
        ("4. No LangChain / agent framework",
         "~350 lines of orchestrator: OpenAI function calling + rules fallback.",
         "Full control, easy to debug, explain line-by-line in an interview.",
         "At scale with 50+ tools, a platform makes sense. For 7 tools, thin orchestration wins."),
        ("5. Dual engine: LLM + rules fallback",
         "Works without API key. Eval harness runs deterministically.",
         "Rules are safety net + test surface, not the production path.",
         "Production is LLM-primary with rules as guardrails on tool parameters."),
        ("6. Identity verification before refunds",
         "verify_customer_identity runs before initiate_refund.",
         "Refunds are fraud-sensitive. Wrong email stops flow before refund record created.",
         "Email match is weak auth for prototype. Production: OTP, session, billing ZIP."),
        ("7. Human handoff with conversation summary",
         "escalate_to_human packages transcript + ticket ID (ESC-XXXXXX).",
         "#1 escalation complaint: repeating yourself. Agents need context immediately.",
         "Today keyword-triggered. Production adds confidence thresholds + sentiment."),
        ("8. Eval harness (30 golden cases)",
         "JSON test cases assert tool calls and reply content.",
         "LLM phrasing drifts; tool selection shouldn't. Shows SE rigor.",
         "Next layer: LLM-as-judge for tone, shadow-mode in staging."),
        ("9. JSON files instead of real APIs",
         "data/orders.json, catalog.json, policies.json as OMS/inventory stand-ins.",
         "Tool interfaces are production-shaped. Swap loaders for HTTP clients.",
         "Boundary is tools/bookly_tools.py. Add retries/timeouts at tool layer."),
        ("10. Streamlit for demo UI",
         "Deploy to Streamlit Cloud in minutes for live presentation.",
         "Not production UI — intentionally deprioritized.",
         "Production = embedded chat + auth + webhook to ticket system."),
    ]
    for title, decision, why, defense in decisions:
        _add_heading(doc, title, 2)
        doc.add_paragraph(f"Decision: {decision}")
        p = doc.add_paragraph(f"Why: {why}")
        p.runs[0].italic = True
        doc.add_paragraph(f"If challenged: {defense}")

    _add_heading(doc, "Tough Questions — Quick Answers", 1)
    _add_table(doc,
        ["Question", "Answer"],
        [
            ["Why not fine-tune?", "Tool calling iterates in hours. Fine-tuning doesn't fix hallucinated order IDs."],
            ["Prompt injection?", "Structured tool params only. Production: allowlisted tools, human approval on high-value refunds."],
            ["Latency?", "Stream responses; cache policies. Refund is multi-turn — accuracy > speed."],
            ["Success metrics?", "Deflection rate, CSAT, refund error rate, tool-call accuracy, time-to-resolution."],
            ["Build next?", "Auth/session pre-fill, OMS integration, confidence escalation, observability."],
        ],
    )

    _add_heading(doc, "Closing (30 seconds)", 1)
    doc.add_paragraph(
        "To summarize: clarify before acting, tools for facts, verify before refunds, "
        "escalate with context, test with golden cases. Next: OMS integrations, authenticated sessions, "
        "confidence-based escalation, production observability. Happy to go deeper on any decision."
    )

    _add_heading(doc, "If Something Breaks", 1)
    _add_table(doc,
        ["Problem", "Recovery"],
        [
            ["Streamlit down", "uvicorn app.main:app --port 8000 or CLI"],
            ["LLM slow/errors", "Works without API key — rules engine"],
            ["Wrong reply", "New conversation + exact demo inputs above"],
            ["Skeptic on quality", "Run python evals/run_evals.py live — 30/30"],
        ],
    )

    doc.save(WORD_OUTPUT)
    print(f"Saved: {WORD_OUTPUT}")


def build_one_pager_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_OUTPUT),
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontSize=22,
        textColor=colors.HexColor("#0F1419"),
        spaceAfter=4,
        leading=26,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#4F8CFF"),
        spaceAfter=10,
        leading=14,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=11,
        textColor=colors.HexColor("#4F8CFF"),
        spaceBefore=8,
        spaceAfter=4,
        leading=13,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=9,
        leading=11.5,
        spaceAfter=3,
    )
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=12,
        bulletIndent=0,
        spaceAfter=2,
    )

    story = []

    story.append(Paragraph("Bookly Customer Support Agent", title_style))
    story.append(Paragraph(
        "Grounded, clarify-first AI support for an online bookstore  ·  "
        "github.com/andrewwchild/booklydemo  ·  booklydemo.streamlit.app",
        subtitle_style,
    ))

    story.append(Paragraph("Thesis", section_style))
    story.append(Paragraph(
        "<b>The best CX agents don't guess — they clarify, then act on real data.</b> "
        "Human in conversation, machine-precise on facts. Transactional queries always use tools; "
        "the LLM never invents order IDs, tracking numbers, or stock levels.",
        body_style,
    ))

    story.append(Paragraph("Architecture", section_style))
    arch_data = [
        ["Layer", "Role"],
        ["Chat UI", "Streamlit (demo) · FastAPI + CLI alternates"],
        ["Orchestrator", "Thin LLM ↔ tool loop · rules fallback · no LangChain"],
        ["Memory", "Per-session slots: order_id, reason, email, book_title"],
        ["Tools (7)", "lookup_order · verify_identity · initiate_refund · check_stock · get_policy · send_password_reset · escalate_to_human"],
        ["Data", "JSON stores (prototype) → OMS / inventory / CRM in production"],
        ["Quality", "30-case golden-set eval harness · 100% pass · CI-ready"],
    ]
    arch_table = Table(arch_data, colWidths=[1.3 * inch, 5.9 * inch])
    arch_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F8CFF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D7DE")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FA")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(arch_table)
    story.append(Spacer(1, 6))

    story.append(Paragraph("Key Design Decisions", section_style))
    for item in [
        "<b>Tools over RAG</b> for transactions — zero hallucination on order/stock/refund data",
        "<b>Clarify-first</b> — ask for order ID, email, reason before acting",
        "<b>Verify before refund</b> — identity check, then eligibility check, then initiate",
        "<b>Human handoff</b> — escalation ticket + full conversation summary",
        "<b>Dual engine</b> — LLM when API key set; deterministic rules for demo + evals",
    ]:
        story.append(Paragraph(f"• {item}", bullet_style))

    story.append(Paragraph("Live Demo Highlights", section_style))
    demo_data = [
        ["Flow", "Input", "Proves"],
        ["Order lookup", '"Where\'s my order?" → ORD-1001', "Clarify-first + grounded tracking"],
        ["Refund", "ORD-1003 · damaged · carol@example.com", "Multi-turn slots + identity verify"],
        ["Edge case", "ORD-1004 · david@example.com", "Verification ≠ approval"],
        ["Escalation", '"Speak to a real person"', "Ticket + transcript handoff"],
        ["Inventory", "Fourth Wing / Atomic Habits", "Live stock via tool, not generation"],
    ]
    demo_table = Table(demo_data, colWidths=[1.1 * inch, 2.8 * inch, 3.3 * inch])
    demo_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F1419")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D7DE")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FA")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(demo_table)
    story.append(Spacer(1, 6))

    story.append(Paragraph("Production Roadmap", section_style))
    roadmap_data = [
        ["Gap Today", "Production Fix"],
        ["JSON data stores", "OMS + inventory + CRM APIs"],
        ["Email-match identity", "Auth session + OTP"],
        ["Keyword escalation", "Confidence threshold + sentiment"],
        ["In-process sessions", "Redis + cross-device continuity"],
    ]
    roadmap_table = Table(roadmap_data, colWidths=[2.5 * inch, 4.7 * inch])
    roadmap_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F8CFF")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D7DE")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F6F8FA")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(roadmap_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph(
        "<b>Andrew Child</b>  ·  Decagon Solutions Engineering Take-Home  ·  "
        "Run demo: <font color='#4F8CFF'>streamlit run app/streamlit_app.py</font>  ·  "
        "Run evals: <font color='#4F8CFF'>python evals/run_evals.py</font>",
        ParagraphStyle("Footer", parent=body_style, fontSize=8, textColor=colors.HexColor("#57606A")),
    ))

    doc.build(story)
    print(f"Saved: {PDF_OUTPUT}")


def main() -> None:
    build_word_doc()
    build_one_pager_pdf()


if __name__ == "__main__":
    main()
